"""
OCR & Document Processing Service
- pdfplumber for native PDFs
- PaddleOCR for scanned PDFs
- python-docx for Word documents
"""
import os
import logging
from typing import Optional
from pathlib import Path

logger = logging.getLogger(__name__)


class OCRService:
    def __init__(self):
        self._paddle_ocr = None

    def _get_paddle_ocr(self):
        if self._paddle_ocr is None:
            try:
                from paddleocr import PaddleOCR
                self._paddle_ocr = PaddleOCR(use_angle_cls=True, lang='en', show_log=False)
                logger.info("✅ PaddleOCR initialized")
            except Exception as e:
                logger.warning(f"PaddleOCR unavailable: {e}")
                self._paddle_ocr = False
        return self._paddle_ocr

    def extract_text_from_pdf(self, file_path: str) -> dict:
        """Extract text from PDF - tries native text first, then OCR"""
        result = {"text": "", "method": "none", "pages": 0, "accuracy_estimate": 0.0}
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                result["pages"] = len(pdf.pages)
                pages_text = []
                for page in pdf.pages:
                    text = page.extract_text() or ""
                    pages_text.append(text)
                full_text = "\n\n".join(pages_text).strip()

            if len(full_text) > 100:  # Good native text extraction
                result["text"] = full_text
                result["method"] = "pdfplumber_native"
                result["accuracy_estimate"] = 0.99
                logger.info(f"PDF native extraction: {len(full_text)} chars from {result['pages']} pages")
                return result

            # Scanned PDF - try OCR
            logger.info("Low text content, attempting OCR...")
            return self._ocr_pdf(file_path, result)

        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            result["error"] = str(e)
            return result

    def _ocr_pdf(self, file_path: str, result: dict) -> dict:
        """OCR fallback for scanned PDFs"""
        try:
            from pdf2image import convert_from_path
            ocr = self._get_paddle_ocr()
            if not ocr:
                result["error"] = "PaddleOCR not available"
                return result

            images = convert_from_path(file_path, dpi=200)
            all_text = []
            confidence_scores = []

            for i, img in enumerate(images):
                import tempfile, numpy as np
                from PIL import Image
                img_array = np.array(img)
                ocr_result = ocr.ocr(img_array, cls=True)
                if ocr_result and ocr_result[0]:
                    page_text = []
                    for line in ocr_result[0]:
                        if line and len(line) >= 2:
                            text = line[1][0]
                            confidence = line[1][1]
                            page_text.append(text)
                            confidence_scores.append(confidence)
                    all_text.append(" ".join(page_text))

            result["text"] = "\n\n".join(all_text)
            result["method"] = "paddleocr"
            result["accuracy_estimate"] = sum(confidence_scores) / len(confidence_scores) if confidence_scores else 0.0
            logger.info(f"OCR complete: {len(result['text'])} chars, confidence: {result['accuracy_estimate']:.2f}")
            return result

        except Exception as e:
            logger.error(f"OCR failed: {e}")
            result["error"] = str(e)
            return result

    def extract_text_from_docx(self, file_path: str) -> dict:
        """Extract text from Word documents"""
        result = {"text": "", "method": "python-docx", "pages": 0, "accuracy_estimate": 0.99}
        try:
            from docx import Document
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Also extract tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_texts.append(row_text)
            
            result["text"] = "\n".join(paragraphs)
            if table_texts:
                result["text"] += "\n\nTABLES:\n" + "\n".join(table_texts)
            result["pages"] = max(1, len(paragraphs) // 30)
            logger.info(f"DOCX extraction: {len(result['text'])} chars")
            return result
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            result["error"] = str(e)
            return result

    def extract_text(self, file_path: str, file_type: str) -> dict:
        """Main entry point - routes to appropriate extractor"""
        file_type = file_type.lower()
        if file_type == "pdf":
            return self.extract_text_from_pdf(file_path)
        elif file_type in ["doc", "docx"]:
            return self.extract_text_from_docx(file_path)
        else:
            return {"text": "", "error": f"Unsupported file type: {file_type}"}


ocr_service = OCRService()
